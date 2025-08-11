"""
📄 Document Parser Service
========================

Serviço especializado para extração de texto de documentos.
Utiliza a biblioteca Marker-PDF conforme especificado no PRD.

Responsabilidades:
- Extrair texto de PDFs com Marker
- Processar documentos DOCX
- Normalizar texto para análise por LLM
"""

import logging
import tempfile
import os
from pathlib import Path
# bytes é um tipo built-in, não precisa ser importado

logger = logging.getLogger(__name__)

class DocumentParser:
    def __init__(self):
        """Inicializa o parser de documentos"""
        self.pdf_converter = None
        try:
            # Tentar diferentes importações do marker (versões diferentes)
            try:
                from marker import PdfConverter, create_model_dict
                logger.info("🔧 Inicializando Marker-PDF converter...")
                self.pdf_converter = PdfConverter(artifact_dict=create_model_dict())
                logger.info("✅ Marker-PDF inicializado com sucesso")
                
            except ImportError:
                # Tentar versão alternativa do marker
                from marker.converters.pdf import PdfConverter
                from marker.models import create_model_dict
                logger.info("🔧 Inicializando Marker-PDF converter (API alternativa)...")
                self.pdf_converter = PdfConverter(artifact_dict=create_model_dict())
                logger.info("✅ Marker-PDF inicializado com sucesso (API alternativa)")
                
        except ImportError as e:
            logger.warning(f"⚠️ Marker-PDF não disponível: {e}")
            logger.warning("📄 Usando extração básica de PDF como fallback")
            
        except Exception as e:
            logger.error(f"❌ Erro ao inicializar Marker-PDF: {e}")
            logger.warning("📄 Usando extração básica de PDF como fallback")

    def parse_document(self, file_content: bytes, filename: str) -> str:
        """
        📄 Extrai texto de documentos
        
        Args:
            file_content: Conteúdo do arquivo em bytes
            filename: Nome do arquivo original
            
        Returns:
            str: Texto extraído e limpo
        """
        file_extension = Path(filename).suffix.lower()
        
        logger.info(f"📄 Processando {filename} ({file_extension})")
        
        try:
            if file_extension == '.pdf':
                return self._parse_pdf(file_content, filename)
            elif file_extension in ['.docx', '.doc']:
                return self._parse_docx(file_content, filename)
            elif file_extension == '.txt':
                return self._parse_txt(file_content)
            else:
                raise ValueError(f"Tipo de arquivo não suportado: {file_extension}")
                
        except Exception as e:
            logger.error(f"❌ Erro ao processar {filename}: {e}")
            raise

    def _parse_pdf(self, file_content: bytes, filename: str) -> str:
        """Extrai texto de PDF usando Marker"""
        if not self.pdf_converter:
            logger.warning("⚠️ Marker-PDF não disponível, usando fallback")
            return self._parse_pdf_fallback(file_content, filename)
        
        try:
            # Criar arquivo temporário
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Processar com Marker
                logger.info("🔍 Extraindo texto com Marker-PDF...")
                result = self.pdf_converter(temp_file_path)
                
                # Extrair texto do resultado
                if hasattr(result, 'text_only'):
                    extracted_text = result.text_only
                elif hasattr(result, 'markdown'):
                    extracted_text = result.markdown
                elif isinstance(result, str):
                    extracted_text = result
                else:
                    extracted_text = str(result)
                
                logger.info(f"✅ Texto extraído: {len(extracted_text)} caracteres")
                return self._clean_text(extracted_text)
                
            finally:
                # Limpar arquivo temporário
                os.unlink(temp_file_path)
                
        except Exception as e:
            logger.error(f"❌ Erro com Marker-PDF: {e}")
            return self._parse_pdf_fallback(file_content, filename)

    def _parse_pdf_fallback(self, file_content: bytes, filename: str) -> str:
        """Fallback para PDFs sem Marker"""
        try:
            import fitz  # PyMuPDF
            
            # Criar arquivo temporário
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = fitz.open(temp_file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                
                logger.info(f"✅ Texto extraído (PyMuPDF): {len(text)} caracteres")
                return self._clean_text(text)
                
            finally:
                os.unlink(temp_file_path)
                
        except ImportError:
            logger.warning("⚠️ PyMuPDF não disponível")
            return f"[ERRO] Não foi possível extrair texto do PDF: {filename}"
        except Exception as e:
            logger.error(f"❌ Erro no fallback PDF: {e}")
            return f"[ERRO] Falha na extração de texto: {str(e)}"

    def _parse_docx(self, file_content: bytes, filename: str) -> str:
        """Extrai texto de documentos DOCX"""
        try:
            from docx import Document
            import io
            
            # Ler DOCX da memória
            doc = Document(io.BytesIO(file_content))
            
            # Extrair texto de todos os parágrafos
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extrair texto de tabelas também
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            extracted_text = '\n'.join(text_parts)
            logger.info(f"✅ Texto extraído (DOCX): {len(extracted_text)} caracteres")
            return self._clean_text(extracted_text)
            
        except ImportError:
            logger.warning("⚠️ python-docx não disponível")
            return f"[PLACEHOLDER] Implementar extração DOCX para: {filename}"
        except Exception as e:
            logger.error(f"❌ Erro ao processar DOCX: {e}")
            return f"[ERRO] Falha na extração DOCX: {str(e)}"

    def _parse_txt(self, file_content: bytes) -> str:
        """Extrai texto de arquivos TXT"""
        try:
            # Tentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    logger.info(f"✅ Texto extraído (TXT, {encoding}): {len(text)} caracteres")
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            # Se falhar todos os encodings
            logger.warning("⚠️ Não foi possível decodificar o arquivo TXT")
            return "[ERRO] Não foi possível decodificar o arquivo de texto"
            
        except Exception as e:
            logger.error(f"❌ Erro ao processar TXT: {e}")
            return f"[ERRO] Falha na extração TXT: {str(e)}"

    def _clean_text(self, text: str) -> str:
        """
        🧹 Limpa e normaliza o texto extraído
        
        Args:
            text: Texto bruto extraído
            
        Returns:
            str: Texto limpo e normalizado
        """
        if not text:
            return ""
        
        # Remover caracteres de controle e normalizar espaços
        import re
        
        # Remover caracteres de controle (exceto quebras de linha)
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalizar quebras de linha
        text = re.sub(r'\r\n|\r', '\n', text)
        
        # Remover linhas vazias excessivas
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remover espaços excessivos
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remover espaços no início e fim de linhas
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()

    def get_text_stats(self, text: str) -> dict:
        """Retorna estatísticas do texto extraído"""
        if not text:
            return {"chars": 0, "words": 0, "lines": 0, "paragraphs": 0}
        
        return {
            "chars": len(text),
            "words": len(text.split()),
            "lines": len(text.split('\n')),
            "paragraphs": len([p for p in text.split('\n\n') if p.strip()])
        }